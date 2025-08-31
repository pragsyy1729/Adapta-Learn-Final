import os
import json
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path
from langchain.schema import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chat_models import init_chat_model
import PyPDF2
import docx
import re

# Configuration
os.environ["GOOGLE_API_KEY"] = "AIzaSyDbIapvce0L6M7G7v-a31UuIRAHaZnSs9k"

class DocumentProcessor:
    def __init__(self):
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")
        self.llm = init_chat_model("gemini-2.5-flash", model_provider="google_genai")
        self.vector_stores = {}  # Cache for vector stores
        self.base_path = Path(__file__).parent.parent.parent.parent  # Project root
        # Don't create vector_store directory since we use department-specific folders

    def get_vector_store_path(self, department_id: str) -> Path:
        """Get the path for a department's vector store."""
        path = self.base_path / f"vector_store_{department_id}"
        print(f"DEBUG: Generated vector store path: {path}")
        return path

    def get_vector_store(self, department_id: str) -> FAISS:
        """Get or create vector store for a department."""
        if department_id in self.vector_stores:
            print(f"DEBUG: Returning cached vector store for {department_id}")
            return self.vector_stores[department_id]

        store_path = self.get_vector_store_path(department_id)
        print(f"DEBUG: Looking for vector store at: {store_path}")
        print(f"DEBUG: Path exists: {store_path.exists()}")

        if store_path.exists():
            try:
                print(f"DEBUG: Loading existing vector store for {department_id}")
                vector_store = FAISS.load_local(
                    str(store_path),
                    self.embeddings,
                    allow_dangerous_deserialization=True
                )
                self.vector_stores[department_id] = vector_store
                print(f"DEBUG: Successfully loaded vector store with {len(vector_store.docstore._dict)} documents")
                return vector_store
            except Exception as e:
                print(f"Error loading existing vector store for {department_id}: {e}")

        # Create empty vector store if none exists
        print(f"DEBUG: Creating new empty vector store for {department_id}")
        vector_store = FAISS.from_texts(
            ["Initial document for department knowledge base"],
            self.embeddings,
            metadatas=[{"type": "system", "department_id": department_id}]
        )
        self.vector_stores[department_id] = vector_store
        return vector_store

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            print(f"Error extracting text from PDF {file_path}: {e}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {e}")
            return ""

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading TXT file {file_path}: {e}")
            return ""

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats."""
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            print(f"Unsupported file format: {file_extension}")
            return ""

    def chunk_text(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """Split text into chunks."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        return text_splitter.split_text(text)

    def extract_key_concepts(self, text: str) -> List[str]:
        """Extract key concepts from text using LLM."""
        try:
            prompt = f"""
            Extract key concepts, topics, and important terms from the following text.
            Return them as a JSON array of strings.

            Text: {text[:2000]}

            Key concepts:
            """

            response = self.llm.invoke(prompt)
            content = response.content.strip()

            # Try to parse as JSON
            try:
                concepts = json.loads(content)
                return concepts if isinstance(concepts, list) else []
            except:
                # Fallback: extract words that look like concepts
                words = re.findall(r'\b[A-Z][a-z]+\b', content)
                return list(set(words))[:10]

        except Exception as e:
            print(f"Error extracting key concepts: {e}")
            return []

    def create_knowledge_graph(self, chunks: List[str], concepts: List[str]) -> Dict[str, Any]:
        """Create a knowledge graph structure from chunks and concepts."""
        nodes = []
        relationships = []

        # Create nodes from concepts
        for concept in concepts:
            node_id = str(uuid.uuid4())
            nodes.append({
                "id": node_id,
                "label": concept,
                "type": "concept",
                "chunks": []
            })

        # Associate chunks with concepts
        for i, chunk in enumerate(chunks):
            chunk_id = f"chunk_{i}"
            relevant_concepts = []

            # Simple relevance check - can be improved with embeddings
            for node in nodes:
                if node["label"].lower() in chunk.lower():
                    relevant_concepts.append(node["id"])
                    node["chunks"].append(chunk_id)

            # Create relationships between concepts that appear together
            for j, concept1 in enumerate(relevant_concepts):
                for concept2 in relevant_concepts[j+1:]:
                    relationships.append({
                        "source": concept1,
                        "target": concept2,
                        "type": "related_to",
                        "chunks": [chunk_id]
                    })

        return {
            "nodes": nodes,
            "relationships": relationships,
            "chunks": [{"id": f"chunk_{i}", "content": chunk} for i, chunk in enumerate(chunks)]
        }

    def process_document(self, file_path: str, department_id: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """Process a document and add it to the department's vector store."""
        try:
            # Extract text
            text = self.extract_text_from_file(file_path)
            if not text.strip():
                print(f"No text extracted from {file_path}")
                return {"success": False, "error": "No text extracted"}

            # Create chunks
            chunks = self.chunk_text(text)

            # Extract key concepts
            concepts = self.extract_key_concepts(text)

            # Create knowledge graph
            knowledge_graph = self.create_knowledge_graph(chunks, concepts)

            # Prepare documents for vector store
            documents = []

            # Add chunks
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "type": "chunk",
                        "chunk_id": f"chunk_{i}",
                        "department_id": department_id,
                        "file_path": file_path,
                        "concepts": concepts,
                        **(metadata or {})
                    }
                )
                documents.append(doc)

            # Add concept nodes
            for node in knowledge_graph["nodes"]:
                doc = Document(
                    page_content=f"Concept: {node['label']}",
                    metadata={
                        "type": "node",
                        "node_id": node["id"],
                        "label": node["label"],
                        "department_id": department_id,
                        "chunks": node["chunks"],
                        "file_path": file_path,
                        **(metadata or {})
                    }
                )
                documents.append(doc)

            # Add relationships
            for rel in knowledge_graph["relationships"]:
                doc = Document(
                    page_content=f"Relationship: {rel['source']} -> {rel['target']}",
                    metadata={
                        "type": "relationship",
                        "source": rel["source"],
                        "target": rel["target"],
                        "relationship_type": rel["type"],
                        "department_id": department_id,
                        "chunks": rel["chunks"],
                        "file_path": file_path,
                        **(metadata or {})
                    }
                )
                documents.append(doc)

            # Get existing vector store
            vector_store = self.get_vector_store(department_id)

            # Add documents to vector store
            vector_store.add_documents(documents)

            # Save vector store
            store_path = self.get_vector_store_path(department_id)
            vector_store.save_local(str(store_path))

            print(f"Successfully processed document {file_path} for department {department_id}")
            print(f"Added {len(documents)} documents to vector store")

            # Return detailed processing information
            return {
                "success": True,
                "total_chunks": len(chunks),
                "concepts_extracted": len(concepts),
                "documents_created": len(documents),
                "chunks": len(chunks),
                "nodes": len(knowledge_graph["nodes"]),
                "relationships": len(knowledge_graph["relationships"])
            }

        except Exception as e:
            print(f"Error processing document {file_path}: {e}")
            return {"success": False, "error": str(e)}

    def search_documents(self, query: str, department_id: str, k: int = 5) -> List[Document]:
        """Search documents in department's vector store."""
        try:
            vector_store = self.get_vector_store(department_id)
            return vector_store.similarity_search(query, k=k)
        except Exception as e:
            print(f"Error searching documents: {e}")
            return []

    def get_department_stats(self, department_id: str) -> Dict[str, Any]:
        """Get statistics about a department's knowledge base."""
        try:
            vector_store = self.get_vector_store(department_id)

            # Get all documents
            all_docs = vector_store.similarity_search("", k=1000)  # Get as many as possible

            stats = {
                "total_documents": len(all_docs),
                "chunks": 0,
                "nodes": 0,
                "relationships": 0,
                "concepts": set()
            }

            for doc in all_docs:
                doc_type = doc.metadata.get("type", "unknown")
                if doc_type == "chunk":
                    stats["chunks"] += 1
                elif doc_type == "node":
                    stats["nodes"] += 1
                    label = doc.metadata.get("label", "")
                    if label:
                        stats["concepts"].add(label)
                elif doc_type == "relationship":
                    stats["relationships"] += 1

            stats["concepts"] = list(stats["concepts"])
            stats["unique_concepts"] = len(stats["concepts"])

            return stats

        except Exception as e:
            print(f"Error getting department stats: {e}")
            return {"error": str(e)}

# Global document processor instance
document_processor = DocumentProcessor()
